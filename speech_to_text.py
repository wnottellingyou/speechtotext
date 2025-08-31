#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
語音轉文字程式
功能：
1. 錄製語音並轉換為文字
2. 上傳音訊檔案並轉換為文字
3. 支援中英文識別
4. 導出為docx或txt格式
5. AI智能摘要分析
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
from datetime import datetime
import tempfile
import time
import json
import re

# 條件導入可能缺少的套件
try:
    import speech_recognition as sr
    SPEECH_RECOGNITION_AVAILABLE = True
except ImportError:
    SPEECH_RECOGNITION_AVAILABLE = False
    sr = None

try:
    import pyaudio
    PYAUDIO_AVAILABLE = True
except ImportError:
    PYAUDIO_AVAILABLE = False

try:
    import wave
    WAVE_AVAILABLE = True
except ImportError:
    WAVE_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import whisper
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

try:
    from pydub import AudioSegment
    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("語音轉文字程式 - 含AI摘要功能")
        self.root.geometry("1200x800")  # 增加窗口大小以容納更多內容
        
        # 檢查必要套件
        self.check_dependencies()
        
        # 初始化變數
        self.is_recording = False
        self.is_paused = False  # 新增暫停狀態
        self.recording_thread = None  # 錄音線程引用
        self.audio_data = None
        self.recognizer = sr.Recognizer() if SPEECH_RECOGNITION_AVAILABLE else None
        self.microphone = sr.Microphone() if SPEECH_RECOGNITION_AVAILABLE and PYAUDIO_AVAILABLE else None
        self.whisper_model = None
        self.recording_segments = []  # 存儲多段錄音
        self.recording_start_time = None
        self.pause_start_time = None  # 暫停開始時間
        self.total_pause_time = 0  # 總暫停時間
        self.temp_files = []  # 追蹤臨時檔案
        self.audio_files = []  # 存儲多個音訊檔案的資訊 [{'path': '', 'name': '', 'order': int}]
        self.current_processing_index = 0  # 目前處理的檔案索引
        self.total_elapsed_time = 0  # 累計時間（秒）
        self.enable_timestamps = False  # 是否啟用時間戳
        
        # AI摘要相關變數
        self.openai_api_key = ""  # OpenAI API密鑰
        self.ai_summary_result = ""  # AI摘要結果
        self.enable_ai_summary = False  # 是否啟用AI摘要
        
        # 載入Whisper模型
        if WHISPER_AVAILABLE:
            self.load_whisper_model()
        
        # 建立UI
        self.create_widgets()
        
        # 調整麥克風
        if SPEECH_RECOGNITION_AVAILABLE and PYAUDIO_AVAILABLE:
            self.adjust_microphone()
    
    def check_dependencies(self):
        """檢查並顯示缺少的套件"""
        missing_packages = []
        
        if not SPEECH_RECOGNITION_AVAILABLE:
            missing_packages.append("speech_recognition")
        if not PYAUDIO_AVAILABLE:
            missing_packages.append("pyaudio")
        if not DOCX_AVAILABLE:
            missing_packages.append("python-docx")
        if not WHISPER_AVAILABLE:
            missing_packages.append("openai-whisper")
        if not PYDUB_AVAILABLE:
            missing_packages.append("pydub")
        if not OPENAI_AVAILABLE:
            missing_packages.append("openai")
        
        if missing_packages:
            message = f"警告：以下套件未安裝，部分功能可能無法使用：\n\n"
            for pkg in missing_packages:
                message += f"• {pkg}\n"
            message += f"\n安裝命令：\n"
            message += f"pip install {' '.join(missing_packages)}\n\n"
            message += f"程式將以有限功能模式運行。"
            
            # 延遲顯示消息，讓GUI完全載入
            self.root.after(1000, lambda: messagebox.showwarning("套件檢查", message))
    
    def load_whisper_model(self):
        """載入Whisper模型"""
        if not WHISPER_AVAILABLE:
            print("Whisper套件未安裝，跳過模型載入")
            return
            
        try:
            print("載入Whisper模型中...")
            self.whisper_model = whisper.load_model("base")
            print("Whisper模型載入完成")
        except Exception as e:
            print(f"載入Whisper模型失敗: {e}")
            self.whisper_model = None
    
    def adjust_microphone(self):
        """調整麥克風"""
        if not SPEECH_RECOGNITION_AVAILABLE or not PYAUDIO_AVAILABLE:
            print("語音識別套件未安裝，跳過麥克風調整")
            return
            
        try:
            with self.microphone as source:
                self.recognizer.adjust_for_ambient_noise(source)
            print("麥克風已調整完成")
        except Exception as e:
            print(f"調整麥克風失敗: {e}")
    
    def update_record_status(self, text):
        """安全地更新錄音狀態"""
        try:
            self.record_status.config(text=text)
        except Exception as e:
            print(f"更新狀態失敗: {e}")
    
    def safe_update_result(self, text):
        """安全地更新結果顯示"""
        try:
            self.text_result.delete(1.0, tk.END)
            self.text_result.insert(tk.END, text)
        except Exception as e:
            print(f"更新結果失敗: {e}")
    
    def safe_file_cleanup(self, file_path, max_retries=3):
        """安全地清理檔案，包含重試機制"""
        import time
        for attempt in range(max_retries):
            try:
                if os.path.exists(file_path):
                    os.unlink(file_path)
                return True
            except PermissionError:
                print(f"檔案清理失敗 (嘗試 {attempt + 1}/{max_retries}): {file_path}")
                time.sleep(0.5)  # 等待0.5秒後重試
            except Exception as e:
                print(f"檔案清理錯誤: {e}")
                break
        return False
    
    def cleanup_temp_files(self):
        """清理所有追蹤的臨時檔案"""
        cleaned_files = []
        failed_files = []
        
        for file_path in self.temp_files[:]:  # 使用副本遍歷
            if self.safe_file_cleanup(file_path):
                cleaned_files.append(file_path)
                self.temp_files.remove(file_path)
            else:
                failed_files.append(file_path)
        
        if cleaned_files:
            print(f"已清理 {len(cleaned_files)} 個臨時檔案")
        if failed_files:
            print(f"無法清理 {len(failed_files)} 個檔案: {failed_files}")
        
        return len(failed_files) == 0
    
    def reset_all(self):
        """重置所有狀態和清理資源"""
        try:
            # 停止錄音
            if self.is_recording:
                self.is_recording = False
            
            # 重置錄音狀態
            self.is_paused = False
            self.recording_thread = None
            self.pause_start_time = None
            self.total_pause_time = 0
            
            # 清空音訊資料
            self.audio_data = None
            self.recording_segments = []
            self.recording_start_time = None
            self.audio_files = []  # 清空音訊檔案列表
            self.current_processing_index = 0
            self.total_elapsed_time = 0  # 重置累計時間
            
            # 重置AI相關狀態
            self.ai_summary_result = ""
            
            # 清理臨時檔案
            self.cleanup_temp_files()
            
            # 重置UI狀態
            self.root.after(0, self.update_record_status, "已重置")
            self.root.after(0, lambda: self.record_button.config(state="normal"))
            self.root.after(0, lambda: self.pause_button.config(state="disabled"))
            self.root.after(0, lambda: self.resume_button.config(state="disabled"))
            self.root.after(0, lambda: self.stop_button.config(state="disabled"))
            self.root.after(0, lambda: self.clear_audio_files())  # 清空檔案列表
            self.root.after(0, lambda: self.text_result.delete(1.0, tk.END))
            self.root.after(0, lambda: self.clear_ai_summary())  # 清空AI摘要
            
            print("系統已重置")
            return True
            
        except Exception as e:
            print(f"重置失敗: {e}")
            return False
    
    def create_widgets(self):
        """建立UI元件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 標題
        title_label = ttk.Label(main_frame, text="語音轉文字程式 - 含AI摘要功能", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=4, pady=(0, 20))
        
        # 建立左右兩欄佈局
        left_frame = ttk.Frame(main_frame)
        left_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 10))
        
        right_frame = ttk.Frame(main_frame)
        right_frame.grid(row=1, column=2, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(10, 0))
        
        # ==================== 左側區域 ====================
        
        # 錄音區域
        record_frame = ttk.LabelFrame(left_frame, text="語音錄製", padding="10")
        record_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # 錄音控制按鈕區域
        button_frame = ttk.Frame(record_frame)
        button_frame.grid(row=0, column=0, columnspan=2, pady=(0, 10))
        
        self.record_button = ttk.Button(button_frame, text="開始錄音", command=self.start_recording)
        self.record_button.grid(row=0, column=0, padx=(0, 5))
        
        self.pause_button = ttk.Button(button_frame, text="暫停錄音", command=self.pause_recording, state="disabled")
        self.pause_button.grid(row=0, column=1, padx=(0, 5))
        
        self.resume_button = ttk.Button(button_frame, text="繼續錄音", command=self.resume_recording, state="disabled")
        self.resume_button.grid(row=0, column=2, padx=(0, 5))
        
        self.stop_button = ttk.Button(button_frame, text="結束錄音", command=self.stop_recording, state="disabled")
        self.stop_button.grid(row=0, column=3)
        
        self.record_status = ttk.Label(record_frame, text="準備錄音")
        self.record_status.grid(row=1, column=0, columnspan=2, pady=(5, 0))
        
        # 錄音設定
        settings_frame = ttk.Frame(record_frame)
        settings_frame.grid(row=2, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Label(settings_frame, text="錄音模式：").grid(row=0, column=0, sticky=tk.W)
        self.record_mode_var = tk.StringVar(value="continuous")
        ttk.Radiobutton(settings_frame, text="連續錄音", variable=self.record_mode_var, value="continuous").grid(row=0, column=1, padx=(5, 10))
        ttk.Radiobutton(settings_frame, text="單句錄音", variable=self.record_mode_var, value="single").grid(row=0, column=2)
        
        # 檔案上傳區域
        upload_frame = ttk.LabelFrame(left_frame, text="音訊檔案管理", padding="10")
        upload_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        # 檔案操作按鈕
        file_buttons_frame = ttk.Frame(upload_frame)
        file_buttons_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.upload_button = ttk.Button(file_buttons_frame, text="新增音訊檔案", command=self.add_audio_files)
        self.upload_button.grid(row=0, column=0, padx=(0, 10))
        
        self.clear_files_button = ttk.Button(file_buttons_frame, text="清空列表", command=self.clear_audio_files)
        self.clear_files_button.grid(row=0, column=1, padx=(0, 10))
        
        self.batch_convert_button = ttk.Button(file_buttons_frame, text="批次轉換", command=self.batch_convert_files)
        self.batch_convert_button.grid(row=0, column=2)
        
        # 檔案列表區域
        list_frame = ttk.Frame(upload_frame)
        list_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(10, 0))
        
        # 檔案列表
        columns = ("順序", "檔案名稱", "路徑")
        self.file_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=6)
        
        # 設定欄位標題
        self.file_tree.heading("順序", text="順序")
        self.file_tree.heading("檔案名稱", text="檔案名稱")
        self.file_tree.heading("路徑", text="檔案路徑")
        
        # 設定欄位寬度
        self.file_tree.column("順序", width=60, anchor=tk.CENTER)
        self.file_tree.column("檔案名稱", width=200)
        self.file_tree.column("路徑", width=250)
        
        # 捲軸
        file_scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=file_scrollbar.set)
        
        # 配置網格
        self.file_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        file_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # 順序調整按鈕
        order_frame = ttk.Frame(upload_frame)
        order_frame.grid(row=2, column=0, columnspan=3, pady=(10, 0))
        
        ttk.Button(order_frame, text="上移", command=self.move_up).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(order_frame, text="下移", command=self.move_down).grid(row=0, column=1, padx=(0, 10))
        ttk.Button(order_frame, text="移除選中", command=self.remove_selected).grid(row=0, column=2)
        
        # 識別引擎選擇
        engine_frame = ttk.LabelFrame(left_frame, text="識別引擎", padding="10")
        engine_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.engine_var = tk.StringVar(value="google")
        ttk.Radiobutton(engine_frame, text="Google (需網路)", variable=self.engine_var, value="google").grid(row=0, column=0, padx=(0, 10))
        ttk.Radiobutton(engine_frame, text="Whisper (離線)", variable=self.engine_var, value="whisper").grid(row=0, column=1)
        
        # 語言選擇
        lang_frame = ttk.LabelFrame(left_frame, text="語言設定", padding="10")
        lang_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.language_var = tk.StringVar(value="zh-TW")
        ttk.Radiobutton(lang_frame, text="中文", variable=self.language_var, value="zh-TW").grid(row=0, column=0, padx=(0, 10))
        ttk.Radiobutton(lang_frame, text="英文", variable=self.language_var, value="en-US").grid(row=0, column=1, padx=(0, 10))
        ttk.Radiobutton(lang_frame, text="自動", variable=self.language_var, value="auto").grid(row=0, column=2)
        
        # 進階選項
        options_frame = ttk.LabelFrame(left_frame, text="進階選項", padding="10")
        options_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # 時間戳選項
        self.timestamp_var = tk.BooleanVar(value=False)
        timestamp_check = ttk.Checkbutton(
            options_frame, 
            text="啟用逐字時間戳", 
            variable=self.timestamp_var,
            command=self.toggle_timestamp_option
        )
        timestamp_check.grid(row=0, column=0, sticky=tk.W)
        
        # 時間戳格式說明
        self.timestamp_info = ttk.Label(
            options_frame, 
            text="格式：[MM:SS] 或 [HH:MM:SS]",
            font=("Arial", 8),
            foreground="gray"
        )
        self.timestamp_info.grid(row=1, column=0, sticky=tk.W, pady=(5, 0))
        
        # AI摘要選項
        self.ai_summary_var = tk.BooleanVar(value=False)
        ai_summary_check = ttk.Checkbutton(
            options_frame,
            text="啟用AI智能摘要",
            variable=self.ai_summary_var,
            command=self.toggle_ai_summary_option
        )
        ai_summary_check.grid(row=2, column=0, sticky=tk.W, pady=(10, 0))
        
        # AI摘要說明
        self.ai_summary_info = ttk.Label(
            options_frame,
            text="分析主題領域、研究方法、視角、結論等",
            font=("Arial", 8),
            foreground="gray"
        )
        self.ai_summary_info.grid(row=3, column=0, sticky=tk.W, pady=(5, 0))
        
        # AI API設定區域
        api_frame = ttk.LabelFrame(left_frame, text="AI API設定", padding="10")
        api_frame.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Label(api_frame, text="OpenAI API Key:").grid(row=0, column=0, sticky=tk.W)
        self.api_key_entry = ttk.Entry(api_frame, width=40, show="*")
        self.api_key_entry.grid(row=0, column=1, padx=(10, 0), sticky=(tk.W, tk.E))
        
        ttk.Button(api_frame, text="測試連接", command=self.test_api_connection).grid(row=0, column=2, padx=(10, 0))
        
        self.api_status_label = ttk.Label(api_frame, text="未設定", foreground="gray")
        self.api_status_label.grid(row=1, column=0, columnspan=3, pady=(5, 0), sticky=tk.W)
        
        # 轉換按鈕
        convert_button = ttk.Button(left_frame, text="開始轉換", command=self.convert_speech_to_text)
        convert_button.grid(row=6, column=0, columnspan=2, pady=(10, 0))
        
        # ==================== 右側區域 ====================
        
        # 結果顯示區域
        result_frame = ttk.LabelFrame(right_frame, text="轉換結果", padding="10")
        result_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        self.text_result = scrolledtext.ScrolledText(result_frame, width=50, height=20)
        self.text_result.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # AI摘要結果區域
        ai_result_frame = ttk.LabelFrame(right_frame, text="AI智能摘要", padding="10")
        ai_result_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        self.ai_result_text = scrolledtext.ScrolledText(ai_result_frame, width=50, height=12)
        self.ai_result_text.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 自定義Prompt區域
        prompt_frame = ttk.LabelFrame(ai_result_frame, text="自定義分析", padding="5")
        prompt_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0))
        
        ttk.Label(prompt_frame, text="自定義提示詞:").grid(row=0, column=0, sticky=tk.W)
        
        self.custom_prompt_text = scrolledtext.ScrolledText(prompt_frame, width=50, height=3)
        self.custom_prompt_text.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 10))
        
        # 預設提示詞按鈕
        preset_frame = ttk.Frame(prompt_frame)
        preset_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Button(preset_frame, text="學術分析", command=lambda: self.load_preset_prompt("academic")).grid(row=0, column=0, padx=(0, 5))
        ttk.Button(preset_frame, text="商業報告", command=lambda: self.load_preset_prompt("business")).grid(row=0, column=1, padx=(0, 5))
        ttk.Button(preset_frame, text="會議紀錄", command=lambda: self.load_preset_prompt("meeting")).grid(row=0, column=2, padx=(0, 5))
        ttk.Button(preset_frame, text="清空", command=self.clear_custom_prompt).grid(row=0, column=3)
        
        # AI摘要控制按鈕
        ai_control_frame = ttk.Frame(ai_result_frame)
        ai_control_frame.grid(row=2, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Button(ai_control_frame, text="標準AI摘要", command=self.generate_ai_summary).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(ai_control_frame, text="自定義分析", command=self.generate_custom_analysis).grid(row=0, column=1, padx=(0, 10))
        ttk.Button(ai_control_frame, text="清空摘要", command=self.clear_ai_summary).grid(row=0, column=2)
        
        # 導出按鈕
        export_frame = ttk.Frame(right_frame)
        export_frame.grid(row=2, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Button(export_frame, text="導出TXT", command=self.export_txt).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(export_frame, text="導出DOCX", command=self.export_docx).grid(row=0, column=1, padx=(0, 10))
        ttk.Button(export_frame, text="導出完整報告", command=self.export_full_report).grid(row=0, column=2, padx=(0, 10))
        
        # 重置按鈕
        reset_button = ttk.Button(export_frame, text="重置系統", command=self.reset_program, style="Reset.TButton")
        reset_button.grid(row=0, column=3)
        
        # 配置樣式
        style = ttk.Style()
        style.configure("Reset.TButton", foreground="red")
        
        # 配置權重 - 這是關鍵改善！
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        main_frame.columnconfigure(0, weight=1)  # 左欄
        main_frame.columnconfigure(1, weight=0)  # 左欄延伸
        main_frame.columnconfigure(2, weight=1)  # 右欄
        main_frame.columnconfigure(3, weight=0)  # 右欄延伸
        main_frame.rowconfigure(1, weight=1)
        
        # 左側權重配置
        left_frame.columnconfigure(0, weight=1)
        left_frame.columnconfigure(1, weight=0)
        left_frame.rowconfigure(1, weight=1)  # 檔案管理區域可擴展
        
        # 右側權重配置
        right_frame.columnconfigure(0, weight=1)
        right_frame.columnconfigure(1, weight=0)
        right_frame.rowconfigure(0, weight=2)  # 轉換結果區域
        right_frame.rowconfigure(1, weight=1)  # AI摘要區域
        
        # 各區域內部權重配置
        record_frame.columnconfigure(0, weight=1)
        upload_frame.columnconfigure(0, weight=1)
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        result_frame.columnconfigure(0, weight=1)
        result_frame.rowconfigure(0, weight=1)
        ai_result_frame.columnconfigure(0, weight=1)
        ai_result_frame.rowconfigure(0, weight=1)
        api_frame.columnconfigure(1, weight=1)
        
        # 自定義prompt區域權重配置
        prompt_frame.columnconfigure(0, weight=1)
        prompt_frame.rowconfigure(1, weight=1)
    
    def start_recording(self):
        """開始錄音"""
        if self.is_recording:
            return
            
        self.is_recording = True
        self.is_paused = False
        self.recording_segments = []
        self.recording_start_time = time.time()
        self.total_pause_time = 0
        
        # 更新按鈕狀態
        self.record_button.config(state="disabled")
        self.pause_button.config(state="normal")
        self.resume_button.config(state="disabled")
        self.stop_button.config(state="normal")
        self.record_status.config(text="正在錄音...")
        
        # 在新線程中錄音
        self.recording_thread = threading.Thread(target=self.record_audio, daemon=True)
        self.recording_thread.start()
    
    def pause_recording(self):
        """暫停錄音"""
        if not self.is_recording or self.is_paused:
            return
            
        self.is_paused = True
        self.pause_start_time = time.time()
        
        # 更新按鈕狀態
        self.pause_button.config(state="disabled")
        self.resume_button.config(state="normal")
        self.record_status.config(text="錄音已暫停")
        
        print("錄音已暫停")
    
    def resume_recording(self):
        """繼續錄音"""
        if not self.is_recording or not self.is_paused:
            return
            
        # 計算暫停時間
        if self.pause_start_time:
            self.total_pause_time += time.time() - self.pause_start_time
            self.pause_start_time = None
        
        self.is_paused = False
        
        # 更新按鈕狀態
        self.pause_button.config(state="normal")
        self.resume_button.config(state="disabled")
        self.record_status.config(text="正在錄音...")
        
        print("錄音已繼續")
    
    def stop_recording(self):
        """結束錄音"""
        if not self.is_recording:
            return
            
        # 如果正在暫停，先計算暫停時間
        if self.is_paused and self.pause_start_time:
            self.total_pause_time += time.time() - self.pause_start_time
            self.pause_start_time = None
        
        self.is_recording = False
        self.is_paused = False
        
        # 更新按鈕狀態
        self.record_button.config(state="normal")
        self.pause_button.config(state="disabled")
        self.resume_button.config(state="disabled")
        self.stop_button.config(state="disabled")
        
        # 等待錄音線程結束
        if self.recording_thread and self.recording_thread.is_alive():
            self.recording_thread.join(timeout=2)
        
        # 計算實際錄音時間（扣除暫停時間）
        if self.recording_start_time:
            total_time = time.time() - self.recording_start_time
            actual_recording_time = total_time - self.total_pause_time
            self.record_status.config(text=f"錄音已結束 (實際錄音: {actual_recording_time:.1f}秒)")
        else:
            self.record_status.config(text="錄音已結束")
        
        print("錄音已結束")
    
    def record_audio(self):
        """錄音函數"""
        try:
            with self.microphone as source:
                self.recognizer.adjust_for_ambient_noise(source, duration=1)
            
            if self.record_mode_var.get() == "continuous":
                # 連續錄音模式
                self.continuous_recording()
            else:
                # 單句錄音模式
                self.single_recording()
                
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("錯誤", f"錄音失敗: {e}"))
            self.root.after(0, self.stop_recording)
    
    def reset_program(self):
        """重置程式的方法，可在UI中調用"""
        try:
            success = self.reset_all()
            if success:
                messagebox.showinfo("重置", "系統已成功重置")
            else:
                messagebox.showwarning("重置", "重置過程中遇到一些問題，請檢查終端輸出")
        except Exception as e:
            print(f"重置程式失敗: {e}")
            messagebox.showerror("錯誤", f"重置失敗: {e}")
    
    def toggle_timestamp_option(self):
        """切換時間戳選項"""
        self.enable_timestamps = self.timestamp_var.get()
        if self.enable_timestamps:
            print("已啟用時間戳功能")
        else:
            print("已關閉時間戳功能")
    
    def toggle_ai_summary_option(self):
        """切換AI摘要選項"""
        self.enable_ai_summary = self.ai_summary_var.get()
        if self.enable_ai_summary:
            print("已啟用AI摘要功能")
            if not self.openai_api_key:
                messagebox.showinfo("提示", "請設定OpenAI API Key以使用AI摘要功能")
        else:
            print("已關閉AI摘要功能")
    
    def test_api_connection(self):
        """測試OpenAI API連接"""
        api_key = self.api_key_entry.get().strip()
        if not api_key:
            messagebox.showwarning("警告", "請輸入OpenAI API Key")
            return
        
        try:
            # 設定API key
            self.openai_api_key = api_key
            
            # 測試連接 - 嘗試導入並測試OpenAI
            try:
                import openai
                
                # 檢查版本並選擇適當的API調用方式
                if hasattr(openai, '__version__') and openai.__version__.startswith('1.'):
                    # OpenAI v1.x
                    client = openai.OpenAI(api_key=api_key)
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": "Hello"}],
                        max_tokens=5
                    )
                else:
                    # OpenAI v0.x (舊版)
                    openai.api_key = api_key
                    response = openai.ChatCompletion.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": "Hello"}],
                        max_tokens=5
                    )
                
                self.api_status_label.config(text="API連接成功", foreground="green")
                messagebox.showinfo("成功", "OpenAI API連接測試成功！")
                
            except ImportError:
                self.api_status_label.config(text="缺少openai套件", foreground="red")
                result = messagebox.askyesno("缺少套件", "缺少openai套件，是否要自動安裝？")
                if result:
                    self.install_openai_package()
            
        except Exception as e:
            self.api_status_label.config(text="API連接失敗", foreground="red")
            messagebox.showerror("錯誤", f"API連接失敗：{str(e)}")
    
    def install_openai_package(self):
        """自動安裝OpenAI套件"""
        try:
            import subprocess
            import sys
            
            # 在新線程中安裝套件
            def install_in_thread():
                try:
                    subprocess.run([sys.executable, "-m", "pip", "install", "openai"], 
                                 check=True, capture_output=True, text=True)
                    self.root.after(0, lambda: messagebox.showinfo("成功", "OpenAI套件安裝成功！請重新測試API連接。"))
                    self.root.after(0, lambda: self.api_status_label.config(text="套件安裝完成", foreground="blue"))
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("錯誤", f"安裝失敗：{str(e)}"))
                    self.root.after(0, lambda: self.api_status_label.config(text="安裝失敗", foreground="red"))
            
            threading.Thread(target=install_in_thread, daemon=True).start()
            self.api_status_label.config(text="正在安裝套件...", foreground="blue")
            
        except Exception as e:
            messagebox.showerror("錯誤", f"安裝過程中出現錯誤：{str(e)}")
    
    def generate_ai_summary(self):
        """生成AI摘要"""
        if not self.enable_ai_summary:
            messagebox.showinfo("提示", "請先啟用AI摘要功能")
            return
        
        if not self.openai_api_key:
            messagebox.showwarning("警告", "請先設定並測試OpenAI API Key")
            return
        
        # 取得轉換結果文字
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有轉換結果可以分析")
            return
        
        # 在新線程中執行AI分析
        threading.Thread(target=self.perform_ai_analysis, args=(content,), daemon=True).start()
        
        # 更新狀態
        self.ai_result_text.delete(1.0, tk.END)
        self.ai_result_text.insert(tk.END, "正在進行AI分析，請稍候...\n")
    
    def perform_ai_analysis(self, content):
        """執行AI分析"""
        try:
            import openai
            
            # 構建提示詞
            prompt = self.build_analysis_prompt(content)
            
            # 檢查OpenAI版本並選擇適當的API調用方式
            if hasattr(openai, '__version__') and openai.__version__.startswith('1.'):
                # OpenAI v1.x
                client = openai.OpenAI(api_key=self.openai_api_key)
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "你是一個專業的內容分析助手，擅長分析演講、講座和會議內容。"},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=2000,
                    temperature=0.3
                )
                analysis_result = response.choices[0].message.content
            else:
                # OpenAI v0.x (舊版)
                openai.api_key = self.openai_api_key
                response = openai.ChatCompletion.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "你是一個專業的內容分析助手，擅長分析演講、講座和會議內容。"},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=2000,
                    temperature=0.3
                )
                analysis_result = response.choices[0].message.content
            
            # 儲存分析結果
            self.ai_summary_result = analysis_result
            
            # 更新UI
            self.root.after(0, self.update_ai_result, analysis_result)
            
        except Exception as e:
            error_message = f"AI分析失敗：{str(e)}"
            print(error_message)
            self.root.after(0, self.update_ai_result, error_message)
    
    def build_analysis_prompt(self, content):
        """構建AI分析提示詞"""
        prompt = f"""
請分析以下演講/講座內容，並提供結構化的摘要分析：

【原始內容】
{content}

【分析要求】
請按照以下格式提供詳細分析：

## 📋 內容概覽
- 內容類型：[講座/演講/會議/討論等]
- 估計時長：[根據內容長度估計]
- 語言風格：[正式/非正式/學術/商業等]

## 🎯 主題領域
- 主要領域：
- 次要領域：
- 關鍵議題：

## 🔬 研究方法
- 研究取向：[理論/實務/混合]
- 分析方式：[定性/定量/案例研究等]
- 論證方式：[歸納/演繹/比較分析等]

## 👁️ 研究視角
- 學術觀點：
- 實務角度：
- 批判思維：

## 💡 核心結論
- 主要發現：
- 重要觀點：
- 建議行動：

## 🔗 關鍵概念
- 重要術語：
- 核心理論：
- 引用來源：

## 📊 內容結構
- 開場方式：
- 論述邏輯：
- 結尾總結：

## 💭 補充說明
- 特色亮點：
- 可能疑問：
- 延伸思考：

請用繁體中文回答，保持專業且易懂的語言風格。
"""
        return prompt
    
    def update_ai_result(self, result):
        """更新AI分析結果顯示"""
        self.ai_result_text.delete(1.0, tk.END)
        self.ai_result_text.insert(tk.END, result)
    
    def clear_ai_summary(self):
        """清空AI摘要"""
        self.ai_result_text.delete(1.0, tk.END)
        self.ai_summary_result = ""
    
    def load_preset_prompt(self, preset_type):
        """載入預設提示詞"""
        presets = {
            "academic": """請以學術研究的角度分析這段文本，重點關注：
1. 研究主題和問題
2. 理論框架和方法論
3. 主要發現和論證
4. 學術貢獻和意義
5. 可能的研究限制
請提供詳細的學術分析報告。""",
            
            "business": """請以商業分析的角度解讀這段內容，包含：
1. 核心商業議題
2. 市場機會和挑戰
3. 策略建議和行動方案
4. 風險評估
5. 預期效益和影響
請提供實用的商業洞察。""",
            
            "meeting": """請整理這段會議內容，包括：
1. 會議重點議題
2. 主要討論內容
3. 決議事項和行動計畫
4. 責任分工
5. 後續追蹤事項
請提供結構化的會議紀錄。"""
        }
        
        if preset_type in presets:
            self.custom_prompt_text.delete(1.0, tk.END)
            self.custom_prompt_text.insert(1.0, presets[preset_type])
    
    def clear_custom_prompt(self):
        """清空自定義提示詞"""
        self.custom_prompt_text.delete(1.0, tk.END)
    
    def generate_custom_analysis(self):
        """使用自定義提示詞生成AI分析"""
        content = self.text_result.get(1.0, tk.END).strip()
        custom_prompt = self.custom_prompt_text.get(1.0, tk.END).strip()
        
        if not content:
            messagebox.showwarning("警告", "沒有轉換結果可以分析")
            return
        
        if not custom_prompt:
            messagebox.showwarning("警告", "請輸入自定義提示詞")
            return
        
        if not self.api_key_entry.get().strip():
            messagebox.showwarning("警告", "請先設定OpenAI API Key")
            return
        
        # 在新線程中執行分析
        threading.Thread(target=self._perform_custom_analysis, args=(content, custom_prompt), daemon=True).start()
        
        # 顯示分析中狀態
        self.ai_result_text.delete(1.0, tk.END)
        self.ai_result_text.insert(tk.END, "🤖 正在使用自定義提示詞進行AI分析...\n\n")
        self.ai_result_text.update()
    
    def _perform_custom_analysis(self, content, custom_prompt):
        """執行自定義AI分析"""
        try:
            api_key = self.api_key_entry.get().strip()
            
            # 檢查 OpenAI 版本並使用相應的 API
            try:
                import openai
                if hasattr(openai, '__version__') and openai.__version__.startswith('0.'):
                    # 舊版本 API (0.x)
                    openai.api_key = api_key
                    
                    response = openai.ChatCompletion.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "你是一個專業的文本分析助手，請根據用戶的要求進行深度分析。"},
                            {"role": "user", "content": f"{custom_prompt}\n\n以下是要分析的文本內容：\n\n{content}"}
                        ],
                        max_tokens=2000,
                        temperature=0.7
                    )
                    
                    analysis_result = response.choices[0].message.content
                    
                else:
                    # 新版本 API (1.x)
                    from openai import OpenAI
                    client = OpenAI(api_key=api_key)
                    
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "你是一個專業的文本分析助手，請根據用戶的要求進行深度分析。"},
                            {"role": "user", "content": f"{custom_prompt}\n\n以下是要分析的文本內容：\n\n{content}"}
                        ],
                        max_tokens=2000,
                        temperature=0.7
                    )
                    
                    analysis_result = response.choices[0].message.content
                
                # 在主線程中更新UI
                self.root.after(0, lambda: self.update_custom_analysis_result(analysis_result))
                
            except Exception as api_error:
                error_msg = f"API 調用失敗: {str(api_error)}"
                self.root.after(0, lambda: self.update_custom_analysis_result(f"❌ {error_msg}"))
                
        except Exception as e:
            error_msg = f"自定義分析失敗: {str(e)}"
            self.root.after(0, lambda: self.update_custom_analysis_result(f"❌ {error_msg}"))
    
    def update_custom_analysis_result(self, result):
        """更新自定義分析結果"""
        self.ai_result_text.delete(1.0, tk.END)
        self.ai_result_text.insert(tk.END, "🎯 自定義AI分析結果\n")
        self.ai_result_text.insert(tk.END, "=" * 50 + "\n\n")
        self.ai_result_text.insert(tk.END, result)
    
    def export_full_report(self):
        """導出完整報告（包含原文和AI摘要）"""
        content = self.text_result.get(1.0, tk.END).strip()
        ai_summary = self.ai_result_text.get(1.0, tk.END).strip()
        
        if not content:
            messagebox.showwarning("警告", "沒有轉換結果可以導出")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="儲存完整報告",
            defaultextension=".docx",
            filetypes=[("Word檔案", "*.docx"), ("文字檔案", "*.txt"), ("所有檔案", "*.*")]
        )
        
        if file_path:
            try:
                if file_path.endswith('.docx'):
                    self.export_full_docx_report(file_path, content, ai_summary)
                else:
                    self.export_full_txt_report(file_path, content, ai_summary)
                messagebox.showinfo("成功", f"完整報告已導出至: {file_path}")
            except Exception as e:
                messagebox.showerror("錯誤", f"導出失敗: {e}")
    
    def export_full_docx_report(self, file_path, content, ai_summary):
        """導出完整DOCX報告"""
        doc = Document()
        
        # 標題
        doc.add_heading('語音轉文字完整分析報告', 0)
        doc.add_paragraph(f'生成時間: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('-' * 80)
        
        # 原始轉換結果
        doc.add_heading('📝 語音轉文字結果', 1)
        doc.add_paragraph(content)
        
        # AI摘要結果
        if ai_summary:
            doc.add_heading('🤖 AI智能摘要分析', 1)
            doc.add_paragraph(ai_summary)
        else:
            doc.add_heading('🤖 AI智能摘要分析', 1)
            doc.add_paragraph('尚未進行AI分析')
        
        doc.save(file_path)
    
    def export_full_txt_report(self, file_path, content, ai_summary):
        """導出完整TXT報告"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("語音轉文字完整分析報告\n")
            f.write(f"生成時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 80 + "\n\n")
            
            f.write("📝 語音轉文字結果\n")
            f.write("-" * 40 + "\n")
            f.write(content + "\n\n")
            
            f.write("🤖 AI智能摘要分析\n")
            f.write("-" * 40 + "\n")
            if ai_summary:
                f.write(ai_summary + "\n")
            else:
                f.write("尚未進行AI分析\n")
    
    def format_timestamp(self, seconds):
        """格式化時間戳"""
        total_seconds = int(seconds)
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        secs = total_seconds % 60
        
        if hours > 0:
            return f"[{hours:02d}:{minutes:02d}:{secs:02d}]"
        else:
            return f"[{minutes:02d}:{secs:02d}]"
    
    def get_audio_duration(self, file_path):
        """取得音檔長度（秒）"""
        try:
            audio = AudioSegment.from_file(file_path)
            return len(audio) / 1000.0  # 轉換為秒
        except Exception as e:
            print(f"無法取得音檔長度 {file_path}: {e}")
            return 0
    
    def add_timestamps_to_text(self, text, start_time):
        """為文字添加時間戳"""
        if not text or not self.enable_timestamps:
            return text
        
        # 將文字分割成句子
        sentences = self.split_into_sentences(text)
        if not sentences:
            return text
        
        # 估算每個句子的時間間隔
        total_duration = self.get_estimated_speech_duration(text)
        if total_duration <= 0:
            total_duration = 60  # 預設估計時間
        
        time_per_sentence = total_duration / len(sentences)
        
        # 為每個句子添加時間戳
        timestamped_sentences = []
        current_time = start_time
        
        for sentence in sentences:
            if sentence.strip():  # 只處理非空句子
                timestamp = self.format_timestamp(current_time)
                timestamped_sentences.append(f"{timestamp} {sentence.strip()}")
                current_time += time_per_sentence
        
        return "\n".join(timestamped_sentences)
    
    def split_into_sentences(self, text):
        """將文字分割成句子"""
        import re
        # 使用正則表達式分割句子（支援中英文）
        sentences = re.split(r'[。！？；.!?;]\s*', text)
        # 過濾空句子
        return [s.strip() for s in sentences if s.strip()]
    
    def get_estimated_speech_duration(self, text):
        """估算語音持續時間（基於文字長度）"""
        # 中文字符數 + 英文單詞數
        chinese_chars = len([c for c in text if '\u4e00' <= c <= '\u9fff'])
        english_words = len([w for w in text.split() if w.isalpha()])
        
        # 估算：中文約每分鐘200字，英文約每分鐘150詞
        chinese_duration = chinese_chars / 200 * 60  # 秒
        english_duration = english_words / 150 * 60  # 秒
        
        return max(chinese_duration + english_duration, 10)  # 最少10秒
    
    def get_whisper_timestamps(self, audio_data, start_time=0):
        """使用Whisper取得精確的時間戳"""
        try:
            if not self.whisper_model:
                raise Exception("Whisper模型未載入")
            
            # 將AudioData轉換為臨時檔案
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
            temp_file_path = temp_file.name
            self.temp_files.append(temp_file_path)
            
            try:
                with open(temp_file_path, "wb") as f:
                    f.write(audio_data.get_wav_data())
                
                # 使用Whisper轉換並取得segment資訊
                language = self.language_var.get()
                if language == "auto":
                    result = self.whisper_model.transcribe(temp_file_path, word_timestamps=True)
                else:
                    lang_code = "zh" if language == "zh-TW" else "en"
                    result = self.whisper_model.transcribe(temp_file_path, language=lang_code, word_timestamps=True)
                
                # 處理segments並添加起始時間偏移
                timestamped_text = []
                for segment in result.get("segments", []):
                    segment_start = start_time + segment["start"]
                    timestamp = self.format_timestamp(segment_start)
                    text = segment["text"].strip()
                    if text:
                        timestamped_text.append(f"{timestamp} {text}")
                
                return "\n".join(timestamped_text)
                
            finally:
                # 臨時檔案會在程式結束或重置時統一清理
                pass
                
        except Exception as e:
            print(f"Whisper時間戳轉換失敗: {e}")
            # 如果Whisper失敗，回退到估算方法
            text = self.perform_recognition(audio_data)
            return self.add_timestamps_to_text(text, start_time)
    
    def load_audio_file_for_whisper(self, file_path):
        """為Whisper載入音檔並轉換為AudioData格式"""
        try:
            # 檢查檔案格式並轉換
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext in ['.mp3', '.mp4', '.m4a', '.flac', '.aac', '.ogg']:
                # 使用pydub轉換為WAV格式
                audio = AudioSegment.from_file(file_path)
                
                # 創建臨時WAV檔案
                temp_wav = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                temp_wav_path = temp_wav.name
                self.temp_files.append(temp_wav_path)
                audio.export(temp_wav_path, format="wav")
                
                # 載入轉換後的檔案
                with sr.AudioFile(temp_wav_path) as source:
                    return self.recognizer.record(source)
            else:
                # 直接載入WAV檔案
                with sr.AudioFile(file_path) as source:
                    return self.recognizer.record(source)
                    
        except Exception as e:
            print(f"載入音檔失敗 {file_path}: {e}")
            return None
    
    def add_audio_files(self):
        """新增多個音訊檔案"""
        file_paths = filedialog.askopenfilenames(
            title="選擇音訊檔案",
            filetypes=[
                ("音訊檔案", "*.wav *.mp3 *.mp4 *.m4a *.flac *.aac *.ogg"),
                ("WAV檔案", "*.wav"),
                ("MP3檔案", "*.mp3"),
                ("MP4檔案", "*.mp4"),
                ("所有檔案", "*.*")
            ]
        )
        
        if file_paths:
            # 新增檔案到列表
            for file_path in file_paths:
                file_info = {
                    'path': file_path,
                    'name': os.path.basename(file_path),
                    'order': len(self.audio_files) + 1
                }
                self.audio_files.append(file_info)
            
            # 更新顯示
            self.update_file_list()
            messagebox.showinfo("成功", f"已新增 {len(file_paths)} 個音訊檔案")
    
    def clear_audio_files(self):
        """清空音訊檔案列表"""
        self.audio_files = []
        self.current_processing_index = 0
        self.update_file_list()
    
    def update_file_list(self):
        """更新檔案列表顯示"""
        # 清空現有項目
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        
        # 新增檔案項目
        for i, file_info in enumerate(self.audio_files):
            self.file_tree.insert('', 'end', values=(
                file_info['order'],
                file_info['name'],
                file_info['path']
            ))
    
    def get_selected_item(self):
        """取得選中的項目索引"""
        selection = self.file_tree.selection()
        if not selection:
            return None
        
        item = selection[0]
        values = self.file_tree.item(item, 'values')
        if not values:
            return None
        
        # 找到對應的檔案索引
        for i, file_info in enumerate(self.audio_files):
            if file_info['name'] == values[1] and file_info['path'] == values[2]:
                return i
        return None
    
    def move_up(self):
        """上移選中檔案"""
        index = self.get_selected_item()
        if index is None or index == 0:
            return
        
        # 交換位置
        self.audio_files[index], self.audio_files[index-1] = \
            self.audio_files[index-1], self.audio_files[index]
        
        # 更新順序號
        self.update_order_numbers()
        self.update_file_list()
        
        # 保持選中狀態
        self.select_item_by_index(index-1)
    
    def move_down(self):
        """下移選中檔案"""
        index = self.get_selected_item()
        if index is None or index == len(self.audio_files) - 1:
            return
        
        # 交換位置
        self.audio_files[index], self.audio_files[index+1] = \
            self.audio_files[index+1], self.audio_files[index]
        
        # 更新順序號
        self.update_order_numbers()
        self.update_file_list()
        
        # 保持選中狀態
        self.select_item_by_index(index+1)
    
    def remove_selected(self):
        """移除選中檔案"""
        index = self.get_selected_item()
        if index is None:
            messagebox.showwarning("警告", "請先選擇要移除的檔案")
            return
        
        # 確認移除
        file_name = self.audio_files[index]['name']
        if messagebox.askyesno("確認", f"確定要移除檔案 '{file_name}' 嗎？"):
            del self.audio_files[index]
            self.update_order_numbers()
            self.update_file_list()
    
    def update_order_numbers(self):
        """更新順序號"""
        for i, file_info in enumerate(self.audio_files):
            file_info['order'] = i + 1
    
    def select_item_by_index(self, index):
        """根據索引選中項目"""
        if 0 <= index < len(self.audio_files):
            items = self.file_tree.get_children()
            if index < len(items):
                self.file_tree.selection_set(items[index])
                self.file_tree.focus(items[index])
    
    def batch_convert_files(self):
        """批次轉換所有檔案"""
        if not self.audio_files:
            messagebox.showwarning("警告", "請先新增音訊檔案")
            return
        
        # 在新線程中處理批次轉換
        threading.Thread(target=self.process_batch_conversion, daemon=True).start()
    
    def process_batch_conversion(self):
        """處理批次轉換"""
        try:
            all_text = []
            total_files = len(self.audio_files)
            self.total_elapsed_time = 0  # 重置累計時間
            
            for i, file_info in enumerate(self.audio_files):
                self.current_processing_index = i
                
                # 更新狀態
                status_text = f"正在處理檔案 {i+1}/{total_files}: {file_info['name']}"
                self.root.after(0, self.update_record_status, status_text)
                
                # 載入並轉換檔案
                text = self.convert_single_file(file_info['path'])
                
                # 取得檔案持續時間並累加
                file_duration = self.get_audio_duration(file_info['path'])
                
                if text:
                    # 新增檔案標題
                    section_title = f"\n{'='*50}\n檔案 {i+1}: {file_info['name']}"
                    if self.enable_timestamps:
                        section_title += f" (起始時間: {self.format_timestamp(self.total_elapsed_time)})"
                    section_title += f"\n{'='*50}\n"
                    
                    # 處理文字和時間戳
                    if self.enable_timestamps:
                        # 如果使用Whisper引擎，嘗試使用精確時間戳
                        if self.engine_var.get() == "whisper":
                            try:
                                # 重新載入檔案並取得精確時間戳
                                audio_data = self.load_audio_file_for_whisper(file_info['path'])
                                if audio_data:
                                    formatted_text = self.get_whisper_timestamps(audio_data, self.total_elapsed_time)
                                else:
                                    formatted_text = self.add_timestamps_to_text(text, self.total_elapsed_time)
                            except Exception as e:
                                print(f"Whisper時間戳失敗，使用估算: {e}")
                                formatted_text = self.add_timestamps_to_text(text, self.total_elapsed_time)
                        else:
                            formatted_text = self.add_timestamps_to_text(text, self.total_elapsed_time)
                        
                        all_text.append(section_title + formatted_text)
                    else:
                        all_text.append(section_title + text)
                else:
                    # 轉換失敗的情況
                    error_text = f"\n{'='*50}\n檔案 {i+1}: {file_info['name']}\n{'='*50}\n[轉換失敗]\n"
                    all_text.append(error_text)
                
                # 累加時間到下一個檔案
                self.total_elapsed_time += file_duration
            
            # 合併所有文字並顯示
            final_text = "\n".join(all_text)
            self.root.after(0, self.safe_update_result, final_text)
            self.root.after(0, self.update_record_status, f"批次轉換完成 ({total_files} 個檔案)")
            
            # 如果啟用AI摘要，自動生成摘要
            if self.enable_ai_summary and self.openai_api_key:
                self.root.after(1000, self.generate_ai_summary)  # 延遲1秒後生成摘要
            
        except Exception as e:
            error_text = f"批次轉換失敗: {e}"
            self.root.after(0, self.update_record_status, error_text)
            print(error_text)
    
    def convert_single_file(self, file_path):
        """轉換單一檔案為文字"""
        try:
            # 載入音訊檔案
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext in ['.mp3', '.mp4', '.m4a', '.flac', '.aac', '.ogg']:
                # 使用pydub轉換為WAV格式
                audio = AudioSegment.from_file(file_path)
                
                # 創建臨時WAV檔案
                temp_wav = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                temp_wav_path = temp_wav.name
                self.temp_files.append(temp_wav_path)
                audio.export(temp_wav_path, format="wav")
                
                # 載入轉換後的檔案
                with sr.AudioFile(temp_wav_path) as source:
                    audio_data = self.recognizer.record(source)
            else:
                # 直接載入WAV檔案
                with sr.AudioFile(file_path) as source:
                    audio_data = self.recognizer.record(source)
            
            # 進行語音識別
            return self.perform_recognition(audio_data)
            
        except Exception as e:
            print(f"轉換檔案 {file_path} 失敗: {e}")
            return None
    
    def continuous_recording(self):
        """連續錄音模式"""
        segment_count = 0
        while self.is_recording:
            # 如果暫停，則等待
            if self.is_paused:
                time.sleep(0.1)
                continue
                
            try:
                with self.microphone as source:
                    # 錄製音訊段落（允許較長的靜音間隔）
                    audio = self.recognizer.listen(
                        source, 
                        timeout=1, 
                        phrase_time_limit=None  # 移除短語時間限制
                    )
                    
                    # 檢查是否還在錄音且未暫停
                    if self.is_recording and not self.is_paused:
                        self.recording_segments.append(audio)
                        segment_count += 1
                        
                        # 更新狀態
                        elapsed_time = time.time() - self.recording_start_time - self.total_pause_time
                        if self.pause_start_time:  # 如果正在暫停中
                            current_pause = time.time() - self.pause_start_time
                            elapsed_time -= current_pause
                        
                        status_text = f"錄音中... {segment_count}段 (有效時間: {elapsed_time:.1f}秒)"
                        self.root.after(0, self.update_record_status, status_text)
                    
            except sr.WaitTimeoutError:
                # 超時時繼續等待，不結束錄音
                continue
            except Exception as e:
                if self.is_recording:  # 只在仍在錄音時顯示錯誤
                    print(f"錄音段落錯誤: {e}")
                continue
        
        # 合併所有錄音段落
        if self.recording_segments:
            self.merge_audio_segments()
    
    def single_recording(self):
        """單句錄音模式"""
        try:
            # 等待直到不是暫停狀態
            while self.is_recording and self.is_paused:
                time.sleep(0.1)
            
            if not self.is_recording:
                return
                
            with self.microphone as source:
                # 單次錄音，較短的超時時間
                audio = self.recognizer.listen(source, timeout=10, phrase_time_limit=10)
                
                # 檢查是否還在錄音且未暫停
                if self.is_recording and not self.is_paused:
                    self.audio_data = audio
                    
                    # 自動停止錄音
                    self.root.after(0, self.stop_recording)
                
        except sr.WaitTimeoutError:
            if self.is_recording:  # 只在仍在錄音時顯示超時
                self.root.after(0, self.update_record_status, "錄音超時")
                self.root.after(0, self.stop_recording)
    
    def merge_audio_segments(self):
        """合併多個錄音段落"""
        try:
            if not self.recording_segments:
                return
            
            if len(self.recording_segments) == 1:
                self.audio_data = self.recording_segments[0]
                self.root.after(0, self.update_record_status, "錄音完成")
                return
            
            # 創建臨時文件來合併音訊
            temp_files = []
            
            # 將每個段落保存為臨時文件
            for i, segment in enumerate(self.recording_segments):
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_segment_{i}.wav")
                try:
                    with open(temp_file.name, "wb") as f:
                        f.write(segment.get_wav_data())
                    temp_files.append(temp_file.name)
                    self.temp_files.append(temp_file.name)  # 追蹤臨時檔案
                except Exception as e:
                    print(f"保存段落 {i} 失敗: {e}")
                    continue
            
            if not temp_files:
                print("沒有有效的音訊段落")
                return
            
            # 使用pydub合併音訊
            combined = AudioSegment.empty()
            for temp_file in temp_files:
                try:
                    segment_audio = AudioSegment.from_wav(temp_file)
                    combined += segment_audio
                except Exception as e:
                    print(f"載入段落失敗: {e}")
                    continue
            
            if len(combined) == 0:
                print("合併後音訊為空")
                return
            
            # 保存合併結果
            final_temp = tempfile.NamedTemporaryFile(delete=False, suffix="_combined.wav")
            final_temp_path = final_temp.name
            self.temp_files.append(final_temp_path)  # 追蹤臨時檔案
            
            try:
                combined.export(final_temp_path, format="wav")
                
                # 載入合併後的音訊
                with sr.AudioFile(final_temp_path) as source:
                    self.audio_data = self.recognizer.record(source)
                
                print(f"已合併 {len(self.recording_segments)} 個錄音段落")
                status_text = f"錄音完成 (合併了{len(self.recording_segments)}段)"
                self.root.after(0, self.update_record_status, status_text)
                
            except Exception as e:
                print(f"合併最終處理失敗: {e}")
                # 如果合併失敗，使用第一個段落
                if self.recording_segments:
                    self.audio_data = self.recording_segments[0]
                    self.root.after(0, self.update_record_status, "錄音完成（使用第一段）")
            
            # 清理段落臨時文件（保留最終合併檔案供後續使用）
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                    if temp_file in self.temp_files:
                        self.temp_files.remove(temp_file)
                except:
                    pass
                    
        except Exception as e:
            print(f"合併音訊失敗: {e}")
            # 如果合併失敗，使用第一個段落
            if self.recording_segments:
                self.audio_data = self.recording_segments[0]
                self.root.after(0, self.update_record_status, "錄音完成（合併失敗，使用第一段）")
    
    def upload_audio_file(self):
        """上傳音訊檔案"""
        file_path = filedialog.askopenfilename(
            title="選擇音訊檔案",
            filetypes=[
                ("音訊檔案", "*.wav *.mp3 *.mp4 *.m4a *.flac *.aac *.ogg"),
                ("WAV檔案", "*.wav"),
                ("MP3檔案", "*.mp3"),
                ("MP4檔案", "*.mp4"),
                ("所有檔案", "*.*")
            ]
        )
        
        if file_path:
            self.file_label.config(text=os.path.basename(file_path))
            
            # 載入音訊檔案
            try:
                # 檢查檔案格式並轉換
                file_ext = os.path.splitext(file_path)[1].lower()
                
                if file_ext in ['.mp3', '.mp4', '.m4a', '.flac', '.aac', '.ogg']:
                    # 使用pydub轉換為WAV格式
                    audio = AudioSegment.from_file(file_path)
                    
                    # 創建臨時WAV檔案
                    temp_wav = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                    temp_wav_path = temp_wav.name
                    self.temp_files.append(temp_wav_path)  # 追蹤臨時檔案
                    audio.export(temp_wav_path, format="wav")
                    
                    # 載入轉換後的檔案
                    with sr.AudioFile(temp_wav_path) as source:
                        self.audio_data = self.recognizer.record(source)
                    
                else:
                    # 直接載入WAV檔案
                    with sr.AudioFile(file_path) as source:
                        self.audio_data = self.recognizer.record(source)
                
                messagebox.showinfo("成功", "音訊檔案載入成功")
                
            except Exception as e:
                messagebox.showerror("錯誤", f"載入音訊檔案失敗: {e}")
                print(f"詳細錯誤信息: {e}")
    
    def convert_speech_to_text(self):
        """轉換語音為文字"""
        # 檢查是否有錄音數據或檔案列表
        if not self.audio_data and not self.audio_files:
            messagebox.showwarning("警告", "請先錄音或新增音訊檔案")
            return
        
        self.text_result.delete(1.0, tk.END)
        
        if self.audio_files:
            # 如果有檔案列表，執行批次轉換
            self.text_result.insert(tk.END, "開始批次轉換...\n")
            self.batch_convert_files()
        elif self.audio_data:
            # 如果只有錄音數據，執行單一轉換
            self.text_result.insert(tk.END, "正在轉換錄音中，請稍候...\n")
            self.root.update()
            threading.Thread(target=self.perform_single_conversion, daemon=True).start()
    
    def perform_single_conversion(self):
        """執行單一音訊轉換"""
        try:
            text = self.perform_recognition(self.audio_data)
            if text:
                # 如果啟用時間戳，添加時間戳
                if self.enable_timestamps:
                    # 如果使用Whisper引擎，使用精確時間戳
                    if self.engine_var.get() == "whisper":
                        try:
                            formatted_text = self.get_whisper_timestamps(self.audio_data, 0)
                        except Exception as e:
                            print(f"Whisper時間戳失敗，使用估算: {e}")
                            formatted_text = self.add_timestamps_to_text(text, 0)
                    else:
                        formatted_text = self.add_timestamps_to_text(text, 0)  # 錄音從0開始
                    
                    self.root.after(0, self.safe_update_result, formatted_text)
                else:
                    self.root.after(0, self.safe_update_result, text)
                
                self.root.after(0, self.update_record_status, "轉換完成")
                
                # 如果啟用AI摘要，自動生成摘要
                if self.enable_ai_summary and self.openai_api_key:
                    self.root.after(1000, self.generate_ai_summary)  # 延遲1秒後生成摘要
                
            else:
                self.root.after(0, self.safe_update_result, "轉換失敗")
                self.root.after(0, self.update_record_status, "轉換失敗")
                
        except Exception as e:
            error_text = f"轉換失敗: {e}"
            self.root.after(0, self.safe_update_result, error_text)
            self.root.after(0, self.update_record_status, error_text)
            print(error_text)
    
    def perform_recognition(self, audio_data):
        """執行語音識別並返回文字"""
        try:
            engine = self.engine_var.get()
            language = self.language_var.get()
            
            if engine == "google":
                # 使用Google語音識別
                if language == "auto":
                    return self.recognizer.recognize_google(audio_data)
                else:
                    return self.recognizer.recognize_google(audio_data, language=language)
            
            elif engine == "whisper":
                # 使用Whisper語音識別
                if not self.whisper_model:
                    raise Exception("Whisper模型未載入")
                
                # 將AudioData轉換為臨時檔案
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                temp_file_path = temp_file.name
                self.temp_files.append(temp_file_path)
                
                try:
                    with open(temp_file_path, "wb") as f:
                        f.write(audio_data.get_wav_data())
                    
                    # 使用Whisper轉換
                    if self.enable_timestamps:
                        # 啟用時間戳時，取得詳細的segment資訊
                        if language == "auto":
                            result = self.whisper_model.transcribe(temp_file_path, word_timestamps=True)
                        else:
                            lang_code = "zh" if language == "zh-TW" else "en"
                            result = self.whisper_model.transcribe(temp_file_path, language=lang_code, word_timestamps=True)
                        
                        # 如果是批次轉換的一部分，返回帶有segment資訊的結果
                        if hasattr(self, '_return_segments') and self._return_segments:
                            return result
                        else:
                            return result["text"]
                    else:
                        # 一般轉換
                        if language == "auto":
                            result = self.whisper_model.transcribe(temp_file_path)
                        else:
                            lang_code = "zh" if language == "zh-TW" else "en"
                            result = self.whisper_model.transcribe(temp_file_path, language=lang_code)
                        
                        return result["text"]
                
                finally:
                    # 清理臨時檔案會在程式結束或重置時統一處理
                    pass
            
            else:
                raise Exception(f"不支援的識別引擎: {engine}")
                
        except Exception as e:
            print(f"語音識別失敗: {e}")
            return None
    
    def update_result(self, text):
        """更新結果顯示（保留用於向後相容）"""
        self.safe_update_result(text)
    
    def export_txt(self):
        """導出為TXT檔案"""
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有內容可以導出")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="儲存TXT檔案",
            defaultextension=".txt",
            filetypes=[("文字檔案", "*.txt"), ("所有檔案", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(f"語音轉文字結果\n")
                    f.write(f"轉換時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write("-" * 50 + "\n")
                    f.write(content)
                messagebox.showinfo("成功", f"已成功導出至: {file_path}")
            except Exception as e:
                messagebox.showerror("錯誤", f"導出失敗: {e}")
    
    def export_docx(self):
        """導出為DOCX檔案"""
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有內容可以導出")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="儲存DOCX檔案",
            defaultextension=".docx",
            filetypes=[("Word檔案", "*.docx"), ("所有檔案", "*.*")]
        )
        
        if file_path:
            try:
                doc = Document()
                doc.add_heading('語音轉文字結果', 0)
                doc.add_paragraph(f'轉換時間: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph('-' * 50)
                doc.add_paragraph(content)
                doc.save(file_path)
                messagebox.showinfo("成功", f"已成功導出至: {file_path}")
            except Exception as e:
                messagebox.showerror("錯誤", f"導出失敗: {e}")

def main():
    """主函數"""
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
