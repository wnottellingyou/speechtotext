#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基礎語音轉文字程式
僅使用內建和基本套件
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
import wave
import struct
from datetime import datetime

try:
    import speech_recognition as sr
    SPEECH_RECOGNITION_AVAILABLE = True
except ImportError:
    SPEECH_RECOGNITION_AVAILABLE = False

class BasicSpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("基礎語音轉文字程式")
        self.root.geometry("700x600")
        
        # 初始化變數
        self.audio_data = None
        
        if SPEECH_RECOGNITION_AVAILABLE:
            self.recognizer = sr.Recognizer()
            try:
                self.microphone = sr.Microphone()
                self.mic_available = True
            except:
                self.mic_available = False
        else:
            self.mic_available = False
        
        # 建立UI
        self.create_widgets()
    
    def create_widgets(self):
        """建立UI元件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 標題
        title_label = ttk.Label(main_frame, text="語音轉文字程式", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        # 狀態顯示
        status_frame = ttk.LabelFrame(main_frame, text="系統狀態", padding="10")
        status_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 15))
        
        if SPEECH_RECOGNITION_AVAILABLE:
            if self.mic_available:
                status_text = "✓ 語音識別可用，麥克風已就緒"
            else:
                status_text = "⚠ 語音識別可用，但麥克風不可用"
        else:
            status_text = "✗ 語音識別套件未安裝"
        
        ttk.Label(status_frame, text=status_text).grid(row=0, column=0, sticky=tk.W)
        
        # 錄音區域（如果可用）
        if SPEECH_RECOGNITION_AVAILABLE and self.mic_available:
            record_frame = ttk.LabelFrame(main_frame, text="語音錄製", padding="10")
            record_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 15))
            
            self.record_button = ttk.Button(record_frame, text="開始錄音", command=self.record_audio)
            self.record_button.grid(row=0, column=0, padx=(0, 10))
            
            self.record_status = ttk.Label(record_frame, text="準備錄音")
            self.record_status.grid(row=0, column=1)
            
            ttk.Label(record_frame, text="提示：點擊錄音後請在10秒內說話", font=("Arial", 8)).grid(row=1, column=0, columnspan=2, pady=(5, 0))
        
        # 檔案上傳區域
        upload_frame = ttk.LabelFrame(main_frame, text="音訊檔案上傳", padding="10")
        upload_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 15))
        
        self.upload_button = ttk.Button(upload_frame, text="選擇WAV檔案", command=self.upload_audio_file)
        self.upload_button.grid(row=0, column=0, padx=(0, 10))
        
        self.file_label = ttk.Label(upload_frame, text="未選擇檔案", foreground="gray")
        self.file_label.grid(row=0, column=1)
        
        ttk.Label(upload_frame, text="支援格式：WAV（建議16kHz, 16bit）", font=("Arial", 8)).grid(row=1, column=0, columnspan=2, pady=(5, 0))
        
        # 語言和引擎選擇
        options_frame = ttk.LabelFrame(main_frame, text="設定選項", padding="10")
        options_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 15))
        
        # 語言選擇
        ttk.Label(options_frame, text="語言：").grid(row=0, column=0, sticky=tk.W, padx=(0, 5))
        self.language_var = tk.StringVar(value="zh-TW")
        lang_combo = ttk.Combobox(options_frame, textvariable=self.language_var, width=15)
        lang_combo['values'] = ('zh-TW', 'en-US', 'ja-JP', 'ko-KR')
        lang_combo.grid(row=0, column=1, sticky=tk.W)
        lang_combo.state(['readonly'])
        
        # 轉換按鈕
        convert_frame = ttk.Frame(main_frame)
        convert_frame.grid(row=5, column=0, columnspan=2, pady=(0, 15))
        
        self.convert_button = ttk.Button(convert_frame, text="開始語音轉文字", command=self.convert_speech_to_text)
        self.convert_button.grid(row=0, column=0, padx=(0, 10))
        
        if not SPEECH_RECOGNITION_AVAILABLE:
            self.convert_button.config(state="disabled")
            ttk.Label(convert_frame, text="請先安裝speech_recognition套件", foreground="red").grid(row=0, column=1)
        
        # 結果顯示區域
        result_frame = ttk.LabelFrame(main_frame, text="轉換結果", padding="10")
        result_frame.grid(row=6, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        
        self.text_result = scrolledtext.ScrolledText(result_frame, width=70, height=15, wrap=tk.WORD)
        self.text_result.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 功能按鈕
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=7, column=0, columnspan=2, pady=(0, 10))
        
        ttk.Button(button_frame, text="清除內容", command=self.clear_text).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(button_frame, text="匯出TXT", command=self.export_txt).grid(row=0, column=1, padx=(0, 10))
        ttk.Button(button_frame, text="匯出DOCX", command=self.export_docx).grid(row=0, column=2)
        
        # 配置權重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(6, weight=1)
        result_frame.columnconfigure(0, weight=1)
        result_frame.rowconfigure(0, weight=1)
    
    def record_audio(self):
        """錄製音訊"""
        if not SPEECH_RECOGNITION_AVAILABLE or not self.mic_available:
            messagebox.showerror("錯誤", "錄音功能不可用")
            return
        
        self.record_status.config(text="正在錄音...請說話")
        self.record_button.config(state="disabled", text="錄音中...")
        
        def record():
            try:
                with self.microphone as source:
                    self.recognizer.adjust_for_ambient_noise(source, duration=1)
                    audio = self.recognizer.listen(source, timeout=15, phrase_time_limit=10)
                    self.audio_data = audio
                    self.root.after(0, self.record_complete)
            except sr.WaitTimeoutError:
                self.root.after(0, self.record_timeout)
            except Exception as e:
                self.root.after(0, lambda: self.record_error(str(e)))
        
        threading.Thread(target=record, daemon=True).start()
    
    def record_complete(self):
        """錄音完成"""
        self.record_status.config(text="錄音完成！")
        self.record_button.config(state="normal", text="開始錄音")
    
    def record_timeout(self):
        """錄音超時"""
        self.record_status.config(text="錄音超時，請重試")
        self.record_button.config(state="normal", text="開始錄音")
    
    def record_error(self, error_msg):
        """錄音錯誤"""
        self.record_status.config(text="錄音失敗")
        self.record_button.config(state="normal", text="開始錄音")
        messagebox.showerror("錯誤", f"錄音失敗：{error_msg}")
    
    def upload_audio_file(self):
        """上傳音訊檔案"""
        if not SPEECH_RECOGNITION_AVAILABLE:
            messagebox.showerror("錯誤", "請先安裝speech_recognition套件")
            return
        
        file_path = filedialog.askopenfilename(
            title="選擇音訊檔案",
            filetypes=[
                ("WAV檔案", "*.wav"),
                ("所有檔案", "*.*")
            ]
        )
        
        if file_path:
            try:
                with sr.AudioFile(file_path) as source:
                    self.audio_data = self.recognizer.record(source)
                
                filename = os.path.basename(file_path)
                self.file_label.config(text=f"已載入：{filename}", foreground="blue")
                messagebox.showinfo("成功", "音訊檔案載入成功！")
                
            except Exception as e:
                self.file_label.config(text="載入失敗", foreground="red")
                messagebox.showerror("錯誤", f"載入音訊檔案失敗：{e}")
    
    def convert_speech_to_text(self):
        """轉換語音為文字"""
        if not SPEECH_RECOGNITION_AVAILABLE:
            messagebox.showerror("錯誤", "請先安裝speech_recognition套件")
            return
        
        if not self.audio_data:
            messagebox.showwarning("警告", "請先錄音或上傳音訊檔案")
            return
        
        # 顯示轉換中狀態
        self.text_result.delete(1.0, tk.END)
        self.text_result.insert(tk.END, "🔄 正在轉換中，請稍候...\n\n")
        self.text_result.insert(tk.END, "提示：首次使用可能需要下載語音模型，請耐心等待。")
        self.convert_button.config(state="disabled", text="轉換中...")
        self.root.update()
        
        def convert():
            try:
                language = self.language_var.get()
                text = self.recognizer.recognize_google(self.audio_data, language=language)
                
                # 格式化結果
                result = f"📝 轉換結果：\n\n{text}\n\n"
                result += f"🌐 使用語言：{language}\n"
                result += f"⏰ 轉換時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                
                self.root.after(0, lambda: self.show_result(result, True))
                
            except sr.UnknownValueError:
                error_msg = "❌ 無法識別語音內容\n\n可能原因：\n• 音訊品質不佳\n• 語言設定錯誤\n• 背景噪音過大"
                self.root.after(0, lambda: self.show_result(error_msg, False))
                
            except sr.RequestError as e:
                error_msg = f"🌐 網路請求失敗：{e}\n\n請檢查：\n• 網路連線是否正常\n• 是否有防火牆阻擋"
                self.root.after(0, lambda: self.show_result(error_msg, False))
                
            except Exception as e:
                error_msg = f"⚠ 轉換失敗：{e}"
                self.root.after(0, lambda: self.show_result(error_msg, False))
        
        threading.Thread(target=convert, daemon=True).start()
    
    def show_result(self, text, success):
        """顯示轉換結果"""
        self.text_result.delete(1.0, tk.END)
        self.text_result.insert(tk.END, text)
        self.convert_button.config(state="normal", text="開始語音轉文字")
        
        if success:
            messagebox.showinfo("成功", "語音轉文字完成！")
    
    def clear_text(self):
        """清除文字內容"""
        self.text_result.delete(1.0, tk.END)
    
    def export_txt(self):
        """匯出為TXT檔案"""
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有內容可以匯出")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="儲存TXT檔案",
            defaultextension=".txt",
            filetypes=[("文字檔案", "*.txt"), ("所有檔案", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                messagebox.showinfo("成功", f"已成功匯出至：\n{file_path}")
            except Exception as e:
                messagebox.showerror("錯誤", f"匯出失敗：{e}")
    
    def export_docx(self):
        """匯出為DOCX檔案"""
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有內容可以匯出")
            return
        
        try:
            from docx import Document
            
            file_path = filedialog.asksaveasfilename(
                title="儲存DOCX檔案",
                defaultextension=".docx",
                filetypes=[("Word檔案", "*.docx"), ("所有檔案", "*.*")]
            )
            
            if file_path:
                doc = Document()
                doc.add_heading('語音轉文字結果', 0)
                doc.add_paragraph(content)
                doc.save(file_path)
                messagebox.showinfo("成功", f"已成功匯出至：\n{file_path}")
                
        except ImportError:
            messagebox.showerror("錯誤", "缺少python-docx套件，無法匯出DOCX格式\n請執行：pip install python-docx")
        except Exception as e:
            messagebox.showerror("錯誤", f"匯出失敗：{e}")

def main():
    """主函數"""
    root = tk.Tk()
    app = BasicSpeechToTextApp(root)
    
    # 顯示使用說明
    if not SPEECH_RECOGNITION_AVAILABLE:
        messagebox.showinfo(
            "安裝提示",
            "請先安裝必要套件：\n\n"
            "pip install SpeechRecognition\n"
            "pip install pyaudio\n"
            "pip install python-docx\n\n"
            "或執行install_speech_packages.bat"
        )
    
    root.mainloop()

if __name__ == "__main__":
    main()
