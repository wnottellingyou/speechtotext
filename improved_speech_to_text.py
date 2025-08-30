#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
改進版語音轉文字程式
新增功能：
1. 支援MP3、MP4等多種音訊格式
2. 改善連續錄音功能，避免停歇被視為結束
3. 更好的錯誤處理和用戶體驗
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import speech_recognition as sr
import pyaudio
import wave
import threading
import os
from datetime import datetime
from docx import Document
import whisper
from pydub import AudioSegment
import tempfile
import time

class ImprovedSpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("改進版語音轉文字程式")
        self.root.geometry("900x700")
        
        # 初始化變數
        self.is_recording = False
        self.audio_data = None
        self.recognizer = sr.Recognizer()
        self.microphone = None
        self.whisper_model = None
        self.recording_segments = []
        self.recording_start_time = None
        
        # 初始化麥克風
        self.init_microphone()
        
        # 載入Whisper模型
        self.load_whisper_model()
        
        # 建立UI
        self.create_widgets()
    
    def init_microphone(self):
        """初始化麥克風"""
        try:
            self.microphone = sr.Microphone()
            with self.microphone as source:
                self.recognizer.adjust_for_ambient_noise(source, duration=1)
            print("麥克風已就緒")
        except Exception as e:
            print(f"麥克風初始化失敗: {e}")
            self.microphone = None
    
    def load_whisper_model(self):
        """載入Whisper模型（在背景執行）"""
        def load_in_background():
            try:
                print("正在載入Whisper模型...")
                self.whisper_model = whisper.load_model("base")
                print("Whisper模型載入完成")
                self.root.after(0, lambda: self.update_status("Whisper模型已就緒"))
            except Exception as e:
                print(f"載入Whisper模型失敗: {e}")
                self.whisper_model = None
                self.root.after(0, lambda: self.update_status("Whisper模型載入失敗"))
        
        threading.Thread(target=load_in_background, daemon=True).start()
    
    def update_status(self, message):
        """更新狀態信息"""
        if hasattr(self, 'status_label'):
            self.status_label.config(text=message)
    
    def create_widgets(self):
        """建立UI元件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 標題
        title_label = ttk.Label(main_frame, text="改進版語音轉文字程式", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 15))
        
        # 狀態顯示
        self.status_label = ttk.Label(main_frame, text="系統初始化中...", foreground="blue")
        self.status_label.grid(row=1, column=0, columnspan=3, pady=(0, 15))
        
        # 錄音區域
        record_frame = ttk.LabelFrame(main_frame, text="語音錄製", padding="15")
        record_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # 錄音控制
        control_frame = ttk.Frame(record_frame)
        control_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E))
        
        self.record_button = ttk.Button(control_frame, text="開始錄音", command=self.toggle_recording)
        self.record_button.grid(row=0, column=0, padx=(0, 15))
        
        self.record_status = ttk.Label(control_frame, text="準備錄音")
        self.record_status.grid(row=0, column=1)
        
        if not self.microphone:
            self.record_button.config(state="disabled")
            self.record_status.config(text="麥克風不可用", foreground="red")
        
        # 錄音模式選擇
        mode_frame = ttk.Frame(record_frame)
        mode_frame.grid(row=1, column=0, columnspan=3, pady=(15, 10), sticky=tk.W)
        
        ttk.Label(mode_frame, text="錄音模式：", font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W)
        
        self.record_mode_var = tk.StringVar(value="continuous")
        ttk.Radiobutton(mode_frame, text="連續錄音（推薦）", variable=self.record_mode_var, value="continuous").grid(row=0, column=1, padx=(10, 15))
        ttk.Radiobutton(mode_frame, text="單句錄音", variable=self.record_mode_var, value="single").grid(row=0, column=2)
        
        # 說明文字
        help_frame = ttk.Frame(record_frame)
        help_frame.grid(row=2, column=0, columnspan=3, pady=(5, 0), sticky=tk.W)
        
        help_text = "連續錄音：可以錄製多句話，中間的停頓不會結束錄音\n單句錄音：錄製一句話後自動停止"
        ttk.Label(help_frame, text=help_text, font=("Arial", 8), foreground="gray").grid(row=0, column=0, sticky=tk.W)
        
        # 檔案上傳區域
        upload_frame = ttk.LabelFrame(main_frame, text="音訊檔案上傳", padding="15")
        upload_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        upload_control_frame = ttk.Frame(upload_frame)
        upload_control_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E))
        
        self.upload_button = ttk.Button(upload_control_frame, text="選擇音訊檔案", command=self.upload_audio_file)
        self.upload_button.grid(row=0, column=0, padx=(0, 15))
        
        self.file_label = ttk.Label(upload_control_frame, text="未選擇檔案", foreground="gray")
        self.file_label.grid(row=0, column=1)
        
        # 支援格式說明
        format_text = "支援格式：WAV, MP3, MP4, M4A, FLAC, AAC, OGG"
        ttk.Label(upload_frame, text=format_text, font=("Arial", 8), foreground="gray").grid(row=1, column=0, columnspan=3, pady=(10, 0), sticky=tk.W)
        
        # 識別引擎選擇
        engine_frame = ttk.LabelFrame(main_frame, text="識別引擎", padding="15")
        engine_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.engine_var = tk.StringVar(value="google")
        ttk.Radiobutton(engine_frame, text="Google API（需網路，速度快）", variable=self.engine_var, value="google").grid(row=0, column=0, sticky=tk.W, padx=(0, 20))
        self.whisper_radio = ttk.Radiobutton(engine_frame, text="Whisper（離線，準確度高）", variable=self.engine_var, value="whisper")
        self.whisper_radio.grid(row=0, column=1, sticky=tk.W)
        
        if not self.whisper_model:
            self.whisper_radio.config(state="disabled")
        
        # 語言選擇
        lang_frame = ttk.LabelFrame(main_frame, text="語言設定", padding="15")
        lang_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.language_var = tk.StringVar(value="zh-TW")
        ttk.Radiobutton(lang_frame, text="中文", variable=self.language_var, value="zh-TW").grid(row=0, column=0, padx=(0, 15))
        ttk.Radiobutton(lang_frame, text="英文", variable=self.language_var, value="en-US").grid(row=0, column=1, padx=(0, 15))
        ttk.Radiobutton(lang_frame, text="自動偵測", variable=self.language_var, value="auto").grid(row=0, column=2)
        
        # 轉換按鈕
        convert_frame = ttk.Frame(main_frame)
        convert_frame.grid(row=6, column=0, columnspan=3, pady=(15, 0))
        
        self.convert_button = ttk.Button(convert_frame, text="開始語音轉文字", command=self.convert_speech_to_text, style="Accent.TButton")
        self.convert_button.grid(row=0, column=0)
        
        # 結果顯示區域
        result_frame = ttk.LabelFrame(main_frame, text="轉換結果", padding="15")
        result_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(15, 0))
        
        self.text_result = scrolledtext.ScrolledText(result_frame, width=80, height=18, wrap=tk.WORD)
        self.text_result.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 功能按鈕
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=8, column=0, columnspan=3, pady=(15, 0))
        
        ttk.Button(button_frame, text="清除內容", command=self.clear_text).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(button_frame, text="匯出TXT", command=self.export_txt).grid(row=0, column=1, padx=(0, 10))
        ttk.Button(button_frame, text="匯出DOCX", command=self.export_docx).grid(row=0, column=2)
        
        # 配置權重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(7, weight=1)
        result_frame.columnconfigure(0, weight=1)
        result_frame.rowconfigure(0, weight=1)
        
        # 初始狀態更新
        self.root.after(1000, lambda: self.update_status("系統就緒"))
    
    def toggle_recording(self):
        """切換錄音狀態"""
        if not self.is_recording:
            self.start_recording()
        else:
            self.stop_recording()
    
    def start_recording(self):
        """開始錄音"""
        if not self.microphone:
            messagebox.showerror("錯誤", "麥克風不可用")
            return
        
        self.is_recording = True
        self.record_button.config(text="停止錄音", style="Warning.TButton")
        self.record_status.config(text="正在錄音...", foreground="red")
        
        # 在新線程中錄音
        threading.Thread(target=self.record_audio, daemon=True).start()
    
    def stop_recording(self):
        """停止錄音"""
        self.is_recording = False
        self.record_button.config(text="開始錄音", style="TButton")
        self.record_status.config(text="錄音已停止", foreground="blue")
    
    def record_audio(self):
        """錄音函數"""
        try:
            # 重置錄音數據
            self.recording_segments = []
            self.recording_start_time = time.time()
            
            with self.microphone as source:
                self.recognizer.adjust_for_ambient_noise(source, duration=1)
            
            if self.record_mode_var.get() == "continuous":
                self.continuous_recording()
            else:
                self.single_recording()
                
        except Exception as e:
            messagebox.showerror("錯誤", f"錄音失敗: {e}")
            self.root.after(0, self.stop_recording)
    
    def continuous_recording(self):
        """連續錄音模式"""
        segment_count = 0
        silence_count = 0
        max_silence = 3  # 最大連續靜音次數
        
        while self.is_recording:
            try:
                with self.microphone as source:
                    # 錄製音訊段落
                    audio = self.recognizer.listen(
                        source, 
                        timeout=1, 
                        phrase_time_limit=5  # 每段最多5秒
                    )
                    self.recording_segments.append(audio)
                    segment_count += 1
                    silence_count = 0  # 重置靜音計數
                    
                    # 更新狀態
                    elapsed_time = time.time() - self.recording_start_time
                    self.root.after(0, lambda s=segment_count, t=elapsed_time: self.record_status.config(
                        text=f"錄音中... {s}段 ({t:.1f}秒)", foreground="red"
                    ))
                    
            except sr.WaitTimeoutError:
                silence_count += 1
                if silence_count >= max_silence and segment_count > 0:
                    # 如果已經有錄音段落且連續靜音，自動停止
                    print(f"檢測到{max_silence}秒靜音，自動停止錄音")
                    self.root.after(0, self.stop_recording)
                    break
                continue
            except Exception as e:
                print(f"錄音段落錯誤: {e}")
                continue
        
        # 合併所有錄音段落
        if self.recording_segments:
            self.root.after(0, lambda: self.record_status.config(text="正在處理錄音...", foreground="blue"))
            self.merge_audio_segments()
        else:
            self.root.after(0, lambda: self.record_status.config(text="未錄製到音訊", foreground="orange"))
    
    def single_recording(self):
        """單句錄音模式"""
        try:
            with self.microphone as source:
                audio = self.recognizer.listen(source, timeout=10, phrase_time_limit=10)
                self.audio_data = audio
                self.root.after(0, self.stop_recording)
                self.root.after(0, lambda: self.record_status.config(text="錄音完成", foreground="green"))
                
        except sr.WaitTimeoutError:
            self.root.after(0, lambda: self.record_status.config(text="錄音超時", foreground="orange"))
            self.root.after(0, self.stop_recording)
    
    def merge_audio_segments(self):
        """合併多個錄音段落"""
        try:
            if not self.recording_segments:
                return
            
            if len(self.recording_segments) == 1:
                self.audio_data = self.recording_segments[0]
                self.root.after(0, lambda: self.record_status.config(text="錄音完成", foreground="green"))
                return
            
            # 創建臨時文件來合併音訊
            temp_files = []
            
            for i, segment in enumerate(self.recording_segments):
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_segment_{i}.wav")
                with open(temp_file.name, "wb") as f:
                    f.write(segment.get_wav_data())
                temp_files.append(temp_file.name)
            
            # 使用pydub合併音訊
            combined = AudioSegment.empty()
            for temp_file in temp_files:
                segment_audio = AudioSegment.from_wav(temp_file)
                combined += segment_audio
                # 在段落之間加入短暫間隔
                combined += AudioSegment.silent(duration=200)  # 200ms間隔
            
            # 保存合併結果
            final_temp = tempfile.NamedTemporaryFile(delete=False, suffix="_combined.wav")
            combined.export(final_temp.name, format="wav")
            
            # 載入合併後的音訊
            with sr.AudioFile(final_temp.name) as source:
                self.audio_data = self.recognizer.record(source)
            
            # 清理臨時文件
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                except:
                    pass
            try:
                os.unlink(final_temp.name)
            except:
                pass
            
            print(f"已合併 {len(self.recording_segments)} 個錄音段落")
            self.root.after(0, lambda: self.record_status.config(
                text=f"錄音完成 (合併了{len(self.recording_segments)}段)", foreground="green"
            ))
            
        except Exception as e:
            print(f"合併音訊失敗: {e}")
            if self.recording_segments:
                self.audio_data = self.recording_segments[0]
                self.root.after(0, lambda: self.record_status.config(text="錄音完成（使用第一段）", foreground="green"))
    
    def upload_audio_file(self):
        """上傳音訊檔案"""
        file_path = filedialog.askopenfilename(
            title="選擇音訊檔案",
            filetypes=[
                ("音訊檔案", "*.wav *.mp3 *.mp4 *.m4a *.flac *.aac *.ogg"),
                ("WAV檔案", "*.wav"),
                ("MP3檔案", "*.mp3"),
                ("MP4檔案", "*.mp4"),
                ("所有檔案", "*.*")
            ]
        )
        
        if file_path:
            self.file_label.config(text="載入中...", foreground="blue")
            
            def load_file():
                try:
                    file_ext = os.path.splitext(file_path)[1].lower()
                    filename = os.path.basename(file_path)
                    
                    if file_ext in ['.mp3', '.mp4', '.m4a', '.flac', '.aac', '.ogg']:
                        # 使用pydub轉換為WAV格式
                        print(f"轉換音訊格式: {file_ext}")
                        audio = AudioSegment.from_file(file_path)
                        
                        # 創建臨時WAV檔案
                        temp_wav = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                        audio.export(temp_wav.name, format="wav")
                        
                        # 載入轉換後的檔案
                        with sr.AudioFile(temp_wav.name) as source:
                            self.audio_data = self.recognizer.record(source)
                        
                        # 清理臨時檔案
                        os.unlink(temp_wav.name)
                        
                    else:
                        # 直接載入WAV檔案
                        with sr.AudioFile(file_path) as source:
                            self.audio_data = self.recognizer.record(source)
                    
                    self.root.after(0, lambda: self.file_label.config(
                        text=f"已載入：{filename}", foreground="green"
                    ))
                    self.root.after(0, lambda: messagebox.showinfo("成功", "音訊檔案載入成功"))
                    
                except Exception as e:
                    self.root.after(0, lambda: self.file_label.config(text="載入失敗", foreground="red"))
                    self.root.after(0, lambda: messagebox.showerror("錯誤", f"載入音訊檔案失敗:\n{e}"))
            
            # 在背景執行檔案載入
            threading.Thread(target=load_file, daemon=True).start()
    
    def convert_speech_to_text(self):
        """轉換語音為文字"""
        if not self.audio_data:
            messagebox.showwarning("警告", "請先錄音或上傳音訊檔案")
            return
        
        self.text_result.delete(1.0, tk.END)
        self.text_result.insert(tk.END, "🔄 正在轉換中，請稍候...\n\n")
        
        engine = self.engine_var.get()
        if engine == "whisper" and not self.whisper_model:
            self.text_result.insert(tk.END, "Whisper模型未載入，切換到Google API...")
            engine = "google"
        
        self.text_result.insert(tk.END, f"使用引擎：{engine.upper()}\n")
        self.text_result.insert(tk.END, f"語言設定：{self.language_var.get()}\n\n")
        
        self.convert_button.config(state="disabled", text="轉換中...")
        self.root.update()
        
        # 在新線程中執行轉換
        threading.Thread(target=self.perform_conversion, daemon=True).start()
    
    def perform_conversion(self):
        """執行轉換"""
        try:
            engine = self.engine_var.get()
            language = self.language_var.get()
            
            if engine == "google":
                # 使用Google語音識別
                if language == "auto":
                    text = self.recognizer.recognize_google(self.audio_data)
                else:
                    text = self.recognizer.recognize_google(self.audio_data, language=language)
            
            elif engine == "whisper":
                # 使用Whisper語音識別
                if not self.whisper_model:
                    raise Exception("Whisper模型未載入")
                
                # 將AudioData轉換為臨時檔案
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                with open(temp_file.name, "wb") as f:
                    f.write(self.audio_data.get_wav_data())
                
                # 使用Whisper轉換
                if language == "auto":
                    result = self.whisper_model.transcribe(temp_file.name)
                elif language == "zh-TW":
                    result = self.whisper_model.transcribe(temp_file.name, language="zh")
                else:
                    result = self.whisper_model.transcribe(temp_file.name, language="en")
                
                text = result["text"]
                
                # 清理臨時檔案
                os.unlink(temp_file.name)
            
            # 在主線程中更新UI
            self.root.after(0, lambda: self.show_result(text, True))
            
        except sr.UnknownValueError:
            error_msg = "❌ 無法識別語音內容\n\n可能原因：\n• 音訊品質不佳\n• 語言設定錯誤\n• 背景噪音過大"
            self.root.after(0, lambda: self.show_result(error_msg, False))
            
        except sr.RequestError as e:
            error_msg = f"🌐 網路請求失敗：{e}\n\n請檢查：\n• 網路連線是否正常\n• 是否有防火牆阻擋"
            self.root.after(0, lambda: self.show_result(error_msg, False))
            
        except Exception as e:
            error_msg = f"⚠ 轉換失敗：{e}"
            self.root.after(0, lambda: self.show_result(error_msg, False))
    
    def show_result(self, text, success):
        """顯示轉換結果"""
        self.text_result.delete(1.0, tk.END)
        
        if success:
            result = f"✅ 轉換成功！\n\n"
            result += f"📝 結果：\n{text}\n\n"
            result += f"⏰ 轉換時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            result += f"🔧 使用引擎：{self.engine_var.get().upper()}\n"
            result += f"🌐 語言設定：{self.language_var.get()}"
        else:
            result = text
        
        self.text_result.insert(tk.END, result)
        self.convert_button.config(state="normal", text="開始語音轉文字")
        
        if success:
            messagebox.showinfo("成功", "語音轉文字轉換完成！")
    
    def clear_text(self):
        """清除文字內容"""
        self.text_result.delete(1.0, tk.END)
    
    def export_txt(self):
        """匯出為TXT檔案"""
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有內容可以匯出")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="儲存TXT檔案",
            defaultextension=".txt",
            filetypes=[("文字檔案", "*.txt"), ("所有檔案", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                messagebox.showinfo("成功", f"已成功匯出至：\n{file_path}")
            except Exception as e:
                messagebox.showerror("錯誤", f"匯出失敗：{e}")
    
    def export_docx(self):
        """匯出為DOCX檔案"""
        content = self.text_result.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("警告", "沒有內容可以匯出")
            return
        
        try:
            file_path = filedialog.asksaveasfilename(
                title="儲存DOCX檔案",
                defaultextension=".docx",
                filetypes=[("Word檔案", "*.docx"), ("所有檔案", "*.*")]
            )
            
            if file_path:
                doc = Document()
                doc.add_heading('語音轉文字結果', 0)
                doc.add_paragraph(content)
                doc.save(file_path)
                messagebox.showinfo("成功", f"已成功匯出至：\n{file_path}")
                
        except ImportError:
            messagebox.showerror("錯誤", "缺少python-docx套件，無法匯出DOCX格式")
        except Exception as e:
            messagebox.showerror("錯誤", f"匯出失敗：{e}")

def main():
    """主函數"""
    root = tk.Tk()
    
    # 設定樣式
    style = ttk.Style()
    style.theme_use('winnative')
    
    app = ImprovedSpeechToTextApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
